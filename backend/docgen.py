"""DOCX builders for tailored résumé + cover letter (python-docx).

Shared by the local API (download buttons) — produces a .docx as bytes.
"""

from __future__ import annotations

import io
import re


def _clean_dates(dates) -> str:
    """Drop placeholder 'None' tokens so missing dates don't render as 'None–Present'."""
    s = str(dates or "").strip()
    for tok in ("None", "none", "null", "N/A", "?"):
        s = s.replace(tok, "")
    s = re.sub(r"^[\s–—\-–—]+|[\s–—\-–—]+$", "", s)  # strip leading/trailing dashes
    return s if any(ch.isalnum() for ch in s) else ""


def resume_to_docx(resume: dict, name: str, contact: str) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(name or "")
    r.bold = True
    r.font.size = Pt(16)
    if contact:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.add_run(contact).font.size = Pt(9)
    if resume.get("headline"):
        hl = doc.add_paragraph()
        hl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = hl.add_run(resume["headline"])
        rr.italic = True
        rr.font.size = Pt(11)

    def section(title: str) -> None:
        p = doc.add_paragraph()
        rs = p.add_run(title.upper())
        rs.bold = True
        rs.font.size = Pt(11)

    if resume.get("summary"):
        section("Summary")
        doc.add_paragraph(resume["summary"])
    if resume.get("skills"):
        section("Skills")
        doc.add_paragraph("  •  ".join(resume["skills"]))
    if resume.get("experience"):
        section("Experience")
        for e in resume["experience"]:
            p = doc.add_paragraph()
            rb = p.add_run(e.get("title", ""))
            rb.bold = True
            if e.get("company"):
                p.add_run(f"  —  {e['company']}")
            dates = _clean_dates(e.get("dates"))
            if dates:
                p.add_run(f"   ({dates})").italic = True
            for b in e.get("bullets") or []:
                doc.add_paragraph(b, style="List Bullet")
    if resume.get("education"):
        section("Education")
        for ed in resume["education"]:
            doc.add_paragraph(ed)
    if resume.get("certifications"):
        section("Certifications")
        doc.add_paragraph("  •  ".join(resume["certifications"]))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_inline(paragraph, text: str) -> None:
    """Render a line with **bold** spans into runs on an existing paragraph."""
    for i, chunk in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        if i % 2 == 1:  # odd chunks were inside ** **
            run.bold = True


def master_resume_to_docx(name: str, contact: str, target_title: str, sections) -> bytes:
    """Render the surgically-tailored master résumé (ordered sections) to a .docx.

    Keeps the master's structure; only SUMMARY/CORE SKILLS/title were edited upstream.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()

    def centered(text: str, *, bold=False, italic=False, size=11):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold, r.italic, r.font.size = bold, italic, Pt(size)

    if name:
        centered(name, bold=True, size=16)
    if target_title:
        centered(target_title, italic=True, size=11)
    if contact:
        centered(contact, size=9)

    for head, body in sections:
        sp = doc.add_paragraph()
        rs = sp.add_run(head.upper())
        rs.bold = True
        rs.font.size = Pt(11)
        for raw in body:
            ln = raw.rstrip()
            s = ln.strip()
            if not s or s == "---":
                continue
            if s.startswith("### "):
                p = doc.add_paragraph()
                r = p.add_run(s[4:].strip())
                r.bold = True
                continue
            if s.startswith("- ") or s.startswith("• "):
                p = doc.add_paragraph(style="List Bullet")
                _add_inline(p, s[2:].strip())
                continue
            p = doc.add_paragraph()
            _add_inline(p, s)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def cover_to_docx(text: str, name: str, contact: str, company: str) -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_paragraph().add_run(name or "").bold = True
    if contact:
        doc.add_paragraph(contact).runs[0].font.size = Pt(9)
    if company:
        doc.add_paragraph(company)
    doc.add_paragraph("")
    for para in re.split(r"\n{2,}", (text or "").strip()):
        doc.add_paragraph(para.strip())
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def slug(s: str) -> str:
    return (re.sub(r"_+", "_", re.sub(r"[^a-z0-9]+", "_", (s or "").lower())).strip("_") or "doc")[:50]
