#!/usr/bin/env python3
"""Daily job-match digest — generated + emailed entirely on GX10 (free local LLM).

The Cloudflare Worker computes the day's matches (deterministic, no LLM) and
stores them. THIS script — run on GX10 from a daily cron — does the expensive,
fragile parts on the local model and native tooling instead of Workers AI:

  1. GET /digest/batch   — pulls each user's matches + inventory + job text
  2. for every match with fit_score >= DOC_THRESHOLD (default 90):
       - Qwen3.6 (:8093) writes a JD-tailored résumé (JSON) + cover letter
       - python-docx builds résumé.docx + cover_letter.docx
  3. emails the digest via Resend with those .docx files ATTACHED

Why GX10 and not the Worker: $0 (no Workers-AI neuron cap or cost), native
python-docx (no Worker-runtime limits), résumé/personal data never leaves the
box, and no dependency on a Cloudflare model that can be deprecated under us.

Required env
------------
  CF_BASE_URL       e.g. https://reverse-ats-ingest.aries-lao.workers.dev
  CF_INGEST_SECRET  same value as the Worker INGEST_SECRET
  RESEND_API_KEY    Resend key (same one the Worker used)

Optional env
------------
  LLM_BASE_URL          default http://localhost:8093/v1
  LLM_MODEL             default qwen3.6-35b
  DOC_THRESHOLD         min fit_score to generate+attach docs (default 90)
  MAX_DOCS_PER_USER     safety cap on doc generations per user (default 25)
  FROM_EMAIL            default "Reverse ATS <reverse-ats@arieslabs.ai>"
  CANDIDATE_NAME        name for the résumé header (default = email local-part)
  CANDIDATE_CONTACT     contact line under the name (phone · location · links)
  REVERSE_ATS_LOG_LEVEL default INFO

Cron (deploy via safe-crontab — see CLAUDE.md):
  30 14 * * * cd /mnt/crucial-x10/projects/reverse-ats && CF_BASE_URL=… \
    CF_INGEST_SECRET=… RESEND_API_KEY=… CANDIDATE_NAME="Aries Lao, MBA" \
    .venv/bin/python scripts/daily_digest_gx10.py \
    >> /mnt/crucial-x10/projects/reverse-ats/logs/daily_digest.log 2>&1
"""

from __future__ import annotations

import base64
import io
import json
import logging
import os
import re
import sys
import time
from typing import Any
from urllib import request as urlrequest
from urllib.error import HTTPError, URLError

APP_URL = "https://reverse-ats.app"


def _setup_logging() -> logging.Logger:
    level = os.environ.get("REVERSE_ATS_LOG_LEVEL", "INFO").upper()
    logging.basicConfig(
        level=getattr(logging, level, logging.INFO),
        format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S",
    )
    return logging.getLogger("reverse-ats.daily-digest")


log = _setup_logging()

# ── LLM prompts (kept in lockstep with cloudflare/src/feed.ts) ────────────────

RESUME_SYSTEM = """You are an expert résumé writer tailoring a candidate's résumé to ONE specific job for ATS systems.

Apply these rules (the candidate's own tailoring playbook):
1. HEADLINE — the job's EXACT title (ATS title-matching is real); closest specific match if generic.
2. SUMMARY — 2-3 sentences for THIS role; end with one sentence mirroring the JD's top 2-3 must-have phrases verbatim (kept truthful). Do NOT state a hard total-years number.
3. SKILLS — REORDER, don't fabricate. JD-emphasized skills/keywords first, JD's exact terms. Spell out an acronym once next to its short form. Add a keyword only if the inventory supports it. No years next to skills.
4. EXPERIENCE — rewrite bullets to foreground experience relevant to this job, quantified where the source gives numbers, action-verb first.
5. NEVER invent skills, employers, titles, dates, or metrics not in the inventory — only reframe. Frame any career transition/sabbatical as deliberate, not a gap.

Output ONLY valid JSON:
{"headline":"","summary":"","skills":["",""],"experience":[{"company":"","title":"","dates":"","bullets":["",""]}],"education":["",""],"certifications":["",""]}
JSON only — no prose, no markdown fences."""

COVER_SYSTEM = (
    "You write concise, specific cover letters. 3 short paragraphs (~300 words), no fluff, no clichés. "
    "Reference real experience from the candidate inventory and real requirements from the job. "
    "Lead with the candidate's strongest matches to the role. Return ONLY the letter body — no greeting "
    "line, no signature, no markdown."
)


def _env(name: str, default: str | None = None, required: bool = False) -> str:
    v = os.environ.get(name, default)
    if required and not v:
        log.error("%s is required", name)
        sys.exit(2)
    return v or ""


def _http_json(method: str, url: str, secret: str, payload: Any = None, timeout: int = 60) -> Any:
    data = json.dumps(payload).encode() if payload is not None else None
    req = urlrequest.Request(url, data=data, method=method)
    req.add_header("Authorization", f"Bearer {secret}")
    if data is not None:
        req.add_header("Content-Type", "application/json")
    with urlrequest.urlopen(req, timeout=timeout) as resp:
        return json.loads(resp.read().decode())


def _llm(llm_base: str, model: str, system: str, user: str, max_tokens: int, timeout: int = 180) -> str:
    payload = {
        "model": model,
        "messages": [{"role": "system", "content": system}, {"role": "user", "content": user}],
        "max_tokens": max_tokens,
        "temperature": 0.3,
    }
    req = urlrequest.Request(
        llm_base.rstrip("/") + "/chat/completions",
        data=json.dumps(payload).encode(),
        method="POST",
    )
    req.add_header("Content-Type", "application/json")
    with urlrequest.urlopen(req, timeout=timeout) as resp:
        body = json.loads(resp.read().decode())
    return (body.get("choices") or [{}])[0].get("message", {}).get("content", "")


def _parse_json_loose(text: str) -> Any:
    if not text:
        return None
    cleaned = text.strip()
    fence = re.search(r"```(?:json)?\s*\n?([\s\S]*?)```", cleaned)
    if fence:
        cleaned = fence.group(1).strip()
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        s, e = cleaned.find("{"), cleaned.rfind("}")
        if 0 <= s < e:
            try:
                return json.loads(cleaned[s : e + 1])
            except json.JSONDecodeError:
                return None
    return None


def _inventory_digest(inv: dict) -> str:
    lines: list[str] = []
    if inv.get("summary"):
        lines.append(f"Summary: {inv['summary']}")
    skills = inv.get("skills") or []
    if skills:
        names = [s.get("name") if isinstance(s, dict) else str(s) for s in skills]
        lines.append("Skills: " + ", ".join(n for n in names if n))
    for e in inv.get("experience") or []:
        if not isinstance(e, dict):
            continue
        lines.append(f"- {e.get('title','')} at {e.get('company','')} ({e.get('start','?')} – {e.get('end') or 'Present'})")
        for h in e.get("highlights") or []:
            lines.append(f"  • {h}")
    edu = inv.get("education") or []
    if edu:
        lines.append("Education: " + "; ".join(
            f"{x.get('degree','')} {x.get('field','')} {x.get('school','')} {x.get('end','')}".strip()
            for x in edu if isinstance(x, dict)
        ))
    certs = inv.get("certifications") or []
    if certs:
        lines.append("Certifications: " + ", ".join(c.get("name", "") for c in certs if isinstance(c, dict)))
    return "\n".join(lines)


# ── python-docx builders ──────────────────────────────────────────────────────

def _resume_docx(resume: dict, name: str, contact: str) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run(name)
    run.bold = True
    run.font.size = Pt(16)
    if contact:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.add_run(contact).font.size = Pt(9)
    if resume.get("headline"):
        hl = doc.add_paragraph()
        hl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = hl.add_run(resume["headline"])
        r.italic = True
        r.font.size = Pt(11)

    def section(title: str) -> None:
        p = doc.add_paragraph()
        r = p.add_run(title.upper())
        r.bold = True
        r.font.size = Pt(11)

    if resume.get("summary"):
        section("Summary")
        doc.add_paragraph(resume["summary"])
    if resume.get("skills"):
        section("Skills")
        doc.add_paragraph("  •  ".join(resume["skills"]))
    if resume.get("experience"):
        section("Experience")
        for e in resume["experience"]:
            p = doc.add_paragraph()
            rb = p.add_run(e.get("title", ""))
            rb.bold = True
            if e.get("company"):
                p.add_run(f"  —  {e['company']}")
            if e.get("dates"):
                p.add_run(f"   ({e['dates']})").italic = True
            for b in e.get("bullets") or []:
                doc.add_paragraph(b, style="List Bullet")
    if resume.get("education"):
        section("Education")
        for ed in resume["education"]:
            doc.add_paragraph(ed)
    if resume.get("certifications"):
        section("Certifications")
        doc.add_paragraph("  •  ".join(resume["certifications"]))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _cover_docx(text: str, name: str, contact: str, company: str) -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_paragraph().add_run(name).bold = True
    if contact:
        doc.add_paragraph(contact).runs[0].font.size = Pt(9)
    if company:
        doc.add_paragraph(company)
    doc.add_paragraph("")
    for para in re.split(r"\n{2,}", text.strip()):
        doc.add_paragraph(para.strip())
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _slug(s: str) -> str:
    return re.sub(r"_+", "_", re.sub(r"[^a-z0-9]+", "_", (s or "").lower())).strip("_")[:50] or "doc"


# ── email (Resend with attachments) ───────────────────────────────────────────

def _send_email(resend_key: str, from_email: str, to: str, subject: str, html: str, attachments: list[dict]) -> bool:
    payload = {"from": from_email, "to": [to], "subject": subject, "html": html}
    if attachments:
        payload["attachments"] = attachments
    req = urlrequest.Request(
        "https://api.resend.com/emails", data=json.dumps(payload).encode(), method="POST"
    )
    req.add_header("Authorization", f"Bearer {resend_key}")
    req.add_header("Content-Type", "application/json")
    try:
        with urlrequest.urlopen(req, timeout=30) as resp:
            return 200 <= resp.status < 300
    except HTTPError as e:
        log.error("resend failed: %s %s", e.code, e.read().decode()[:200])
        return False
    except URLError as e:
        log.error("resend error: %s", e)
        return False


def _esc(s: str) -> str:
    return (str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;"))


def _build_html(date: str, matches: list[dict], doc_job_ids: set[str]) -> str:
    rows = []
    for m in matches:
        attached = m["job_id"] in doc_job_ids
        strengths = (
            f'<div style="margin-top:6px"><span style="color:#16a34a;font-weight:600;font-size:12px">Strengths:</span> '
            f'<span style="color:#444;font-size:12px">{_esc(", ".join(m.get("strengths", [])))}</span></div>'
            if m.get("strengths") else ""
        )
        gaps = (
            f'<div style="margin-top:2px"><span style="color:#ca8a04;font-weight:600;font-size:12px">Gaps:</span> '
            f'<span style="color:#444;font-size:12px">{_esc(", ".join(m.get("gaps", [])))}</span></div>'
            if m.get("gaps") else ""
        )
        badge_docs = (
            '<span style="display:inline-block;margin-left:8px;padding:2px 8px;border-radius:6px;background:#dbeafe;color:#2563eb;font-size:11px;font-weight:600">📎 résumé + cover letter attached</span>'
            if attached else ""
        )
        rows.append(
            f'<tr><td style="padding:14px 0;border-bottom:1px solid #eee">'
            f'<div><span style="font-weight:600;font-size:15px;color:#111">{_esc(m["title"])}</span>'
            f'<span style="display:inline-block;margin-left:8px;padding:2px 8px;border-radius:6px;background:#dcfce7;color:#16a34a;font-size:12px;font-weight:600">{m["fit_score"]}% fit · {m["coverage_pct"]}% skills</span>{badge_docs}</div>'
            f'<div style="color:#666;font-size:13px;margin-top:2px">{_esc(m["company"])}'
            f'{(" · " + _esc(m["location"])) if m.get("location") else ""}</div>'
            f"{strengths}{gaps}"
            f'<div style="margin-top:8px"><a href="{_esc(m["url"])}" style="font-size:13px;color:#2563eb;text-decoration:none">Open posting & apply →</a></div>'
            f"</td></tr>"
        )
    return (
        '<div style="font-family:-apple-system,Segoe UI,Roboto,Helvetica,Arial,sans-serif;max-width:640px;margin:0 auto;padding:24px">'
        f'<h1 style="font-size:20px;color:#111;margin:0 0 4px">Your {len(matches)} job matches</h1>'
        f'<p style="color:#666;font-size:14px;margin:0 0 16px">{date} · ranked by fit to your skills & experience. '
        f"Tailored résumé + cover letter attached for your strongest fits.</p>"
        f'<table style="width:100%;border-collapse:collapse">{"".join(rows)}</table>'
        f'<div style="margin-top:20px"><a href="{APP_URL}/app/matches" style="display:inline-block;background:#2563eb;color:#fff;font-size:14px;font-weight:600;text-decoration:none;padding:10px 18px;border-radius:8px">View all matches in Reverse ATS</a></div>'
        "</div>"
    )


# ── main ──────────────────────────────────────────────────────────────────────

def main() -> int:
    base = _env("CF_BASE_URL", required=True).rstrip("/")
    secret = _env("CF_INGEST_SECRET", required=True)
    resend_key = _env("RESEND_API_KEY", required=True)
    llm_base = _env("LLM_BASE_URL", "http://localhost:8093/v1")
    model = _env("LLM_MODEL", "qwen3.6-35b")
    threshold = int(_env("DOC_THRESHOLD", "90"))
    max_docs = int(_env("MAX_DOCS_PER_USER", "25"))
    from_email = _env("FROM_EMAIL", "Reverse ATS <reverse-ats@arieslabs.ai>")

    try:
        batch = _http_json("GET", f"{base}/digest/batch", secret)
    except (HTTPError, URLError) as e:
        log.error("failed to fetch digest batch: %s", e)
        return 1
    users = batch.get("users", []) if isinstance(batch, dict) else []
    date = batch.get("date", "")
    if not users:
        log.info("no users with matches today — nothing to send")
        return 0

    for u in users:
        email = u.get("email")
        inv = u.get("inventory")
        matches = u.get("matches", [])
        if not email or not inv or not matches:
            log.info("skip user %s (email=%s, inv=%s, matches=%d)", u.get("user_id"), bool(email), bool(inv), len(matches))
            continue

        name = _env("CANDIDATE_NAME", email.split("@")[0])
        contact = _env("CANDIDATE_CONTACT", email)
        inv_digest = _inventory_digest(inv)

        attachments: list[dict] = []
        doc_job_ids: set[str] = set()
        to_doc = [m for m in matches if m.get("fit_score", 0) >= threshold][:max_docs]
        log.info("user %s: %d matches, generating docs for %d (≥%d%%)", email, len(matches), len(to_doc), threshold)

        for m in to_doc:
            try:
                job_block = (
                    f"{m['title']} at {m['company']}"
                    + (f" ({m['location']})" if m.get("location") else "")
                    + "\n"
                    + (f"Required skills: {', '.join(m.get('required_skills', []))}\n" if m.get("required_skills") else "")
                    + (f"Nice-to-have: {', '.join(m.get('nice_skills', []))}\n" if m.get("nice_skills") else "")
                    + "\n"
                    + (m.get("description") or "")
                )
                # Tailored résumé
                rj = _parse_json_loose(_llm(
                    llm_base, model, RESUME_SYSTEM,
                    f"## Candidate inventory\n{inv_digest[:5000]}\n\n## Target job\n{job_block}\n\nProduce the tailored résumé JSON.",
                    max_tokens=2500,
                ))
                # Cover letter
                cl = _llm(
                    llm_base, model, COVER_SYSTEM,
                    f"Candidate inventory:\n{inv_digest[:4000]}\n\nJob:\n{job_block}\n\nWrite a 3-paragraph cover letter (~300 words).",
                    max_tokens=800,
                ).strip()

                if isinstance(rj, dict):
                    rbytes = _resume_docx(rj, name, contact)
                    attachments.append({
                        "filename": f"resume_{_slug(m['company'])}_{_slug(m['title'])}.docx",
                        "content": base64.b64encode(rbytes).decode(),
                    })
                if cl:
                    cbytes = _cover_docx(cl, name, contact, m["company"])
                    attachments.append({
                        "filename": f"cover_{_slug(m['company'])}_{_slug(m['title'])}.docx",
                        "content": base64.b64encode(cbytes).decode(),
                    })
                doc_job_ids.add(m["job_id"])
            except Exception as e:  # never let one job kill the digest
                log.warning("doc gen failed for %s: %s", m.get("job_id"), e)

        html = _build_html(date, matches, doc_job_ids)
        subject = f"{len(matches)} job matches for you — {date}"
        ok = _send_email(resend_key, from_email, email, subject, html, attachments)
        log.info("user %s: emailed=%s, %d attachments", email, ok, len(attachments))

    return 0


if __name__ == "__main__":
    sys.exit(main())
